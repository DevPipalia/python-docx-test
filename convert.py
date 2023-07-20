# from jinja2 import Environment, FileSystemLoader
# import pdfkit

# # Data to be included in the template
# data = {
#     'title': 'My PDF Document',
#     'heading': 'Hello, PDF!',
#     'content': 'This is some content for the PDF file.',
# }

# # Path to the directory containing the HTML template
# template_dir = 'c:\\Users\\meha\\OneDrive\\Desktop\\jinja2'
# file_loader = FileSystemLoader(template_dir)
# env = Environment(loader=file_loader)

# # Load the Jinja2 template
# template = env.get_template('sample.html')

# # Render the template with data
# rendered_html = template.render(data)

# # Configuration for pdfkit (set the path to wkhtmltopdf if not in PATH)
# config = pdfkit.configuration(wkhtmltopdf="C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe")

# # Convert HTML to PDF
# output_pdf = 'sample.pdf'
# pdfkit.from_string(rendered_html, output_pdf, configuration=config)

# print(f'PDF successfully generated: {output_pdf}'




import os
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches,Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from html.parser import HTMLParser

# Define your mock CV data here (or import it from another file)
mockCVData = [
    {
        "name": "Fiona Wenhan Zhao",
        "email": "wenhanzhao8890@gmail.com",
        "socialMedia": {
            "linkedin": "fiona@linkedin",
            "twitter": "fiona@twitter",
            "github": "fiona@github"
        },
        "professionalExperience": [
            {
                "title": "Founder",
                "company": "UNIQUE BUNNY",
                "location": "Winnipeg, Canada",
                "startDate": "2014-01-01",
                "endDate": "Present",
                "points": [
                    "Founder and GM of the largest chain boutique in Manitoba that specializes in Japanese & Korean beauty and lifestyle products",
                    "Managed 3 brick-n-mortars and online store with $5Mn+ GMV and $1Mn+ annual revenue & $1.2M free cash flow in 2021",
                    "Created an inventory of X+ products ranging from X categories resulting in a YoY revenue growth of X%",
                    "Improved the customer retention rate by X% by supervising 15 store staff and developing customer service training manuals, teaching product features and selling points",
                    "Performed inventory analysis and improved stock-forecasting mechanism by X% by communicating with vendors, couriers, and Canadian Border Services Agency to ensure on-time, complete delivery of products",
                    "Conducted product-mix optimization drives to analyze consumer behavior and accordingly founded X best-selling products",
                    "Collaborated with X+ marketing firms to run online advertising and in-store marketing by allocating a total budget of X$",
                    "Led the digital transformation of the company by designing and launching the official website that has X MAU",
                    "Managed the company’s social media presence across X platforms by actively posting promotions, blogs, and new products; Accumulated 15k+ followers across multiple platforms"
                ]
            },
            {
                "title": "Boarding Advisor",
                "company": "ST. JOHNS - RAVENSCOURT SCHOOL",
                "location": "Winnipeg, Canada",
                "startDate": "2016-01-01",
                "endDate": "2020-12-31",
                "points": [
                    "Designed & executed efficient study programs; Improved student results by X%",
                    "Mentored 30+ international boarding students, providing each student with peer mentorship sessions to help students adjust to the boarding school environment and improve their academic and social performances",
                    "Planned and executed X stimulating programs and activities, connecting students to the Winnipeg community at large and providing students with a deeper understanding of the Canadian culture"
                ]
            },
            {
                "title": "Counter Manager",
                "company": "HUDSON’S BAY COMPANY",
                "location": "Winnipeg, Canada",
                "startDate": "2013-01-01",
                "endDate": "2014-12-31",
                "points": [
                    "Managed the Clarins Paris counter at the Hudson’s Bay Company – Winnipeg flagship, achieving 30% revenue increase",
                    "Awarded as the Top Sales Associate of the Month – Three times",
                    "Created a client & store management SOP that enhanced the customer experience by offering professional consultations to X+ customers; Efforts yielded strong customer satisfaction, earning recognition from Clarins HQ",
                    "Built a clientele of X+ customers by promoting the products on social media platforms"
                ]
            }
        ],
        "skills": [
            "Digital Marketing",
            "Inventory Management",
            "Customer Service",
            "Data Analysis",
            "Social Media Management"
        ],
        "education": [
            {
                "degree": "Bachelor of Arts",
                "major": "Women and Gender Studies",
                "university": "UNIVERSITY OF WINNIPEG",
                "location": "Winnipeg, Canada",
                "gpa":3.0,
                "startDate":2019,
                "endDate":2022
            }
        ],
        "languages": ["English", "Mandarin"],
        "certifications": [],
        "interests": [
            "Entrepreneurship",
            "Fashion",
            "Blogging"
        ]
    }
]

# Load the Jinja environment and specify the directory containing the template file
env = Environment(loader=FileSystemLoader('.'))
template = env.get_template('resume_template.html')

def strip_tags(html):
    class MLStripper(HTMLParser):
        def __init__(self):
            super().__init__()
            self.reset()
            self.fed = []

        def handle_data(self, d):
            self.fed.append(d)

        def get_data(self):
            return ''.join(self.fed)

    stripper = MLStripper()
    stripper.feed(html)
    return stripper.get_data()

# Function to create a Word document from the HTML content
def create_word_document(cv_data,html_content, output_file):
    plain_text = strip_tags(html_content)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(0.3)  # Adjust the top margin in centimeters
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(cv_data['name'])

    # Set the name font to bold
    name_run.bold = True

    # Set the font size for the name (e.g., set to 16 points)
    name_run.font.size = Pt(16)

    # Center align the name paragraph
    name_paragraph.alignment = 1
    name_paragraph_format = name_paragraph.paragraph_format
    name_paragraph_format.space_after = Pt(1)

    contact_info_paragraph = document.add_paragraph()
    contact_info_paragraph.add_run(cv_data['email'])
    social_media = cv_data.get('socialMedia', {})
    if social_media.get('linkedin') or social_media.get('twitter') or social_media.get('github'):
        contact_info_paragraph.add_run(" | ")
        if social_media.get('linkedin'):
            contact_info_paragraph.add_run(social_media['linkedin'])
            if social_media.get('twitter') or social_media.get('github'):
                contact_info_paragraph.add_run(" | ")
        if social_media.get('twitter'):
            contact_info_paragraph.add_run(social_media['twitter'])
            if social_media.get('github'):
                contact_info_paragraph.add_run(" | ")
        if social_media.get('github'):
            contact_info_paragraph.add_run(social_media['github'])
    contact_info_paragraph.alignment = 1
    contact_info_paragraph_format = contact_info_paragraph.paragraph_format
    contact_info_paragraph_format.space_after = Pt(1)


    # Adding the "Education" title paragraph
    education_title_paragraph = document.add_paragraph()
    education_title_run = education_title_paragraph.add_run("Education")

    # Set the title font to bold
    education_title_run.bold = True

    # Set the font size for the title (e.g., set to 14 points)
    education_title_run.font.size = Pt(12)

    # Center align the "Education" title paragraph
    education_title_paragraph.alignment = 1
    education_title_paragraph_format=education_title_paragraph.paragraph_format
    education_title_paragraph_format.space_after=Pt(0)


    line_style = document.styles.add_style('HorizontalLine', WD_STYLE_TYPE.PARAGRAPH)
    line_style.font.underline = True

    # Add a paragraph with an underscore character and apply the custom style
    line_paragraph = document.add_paragraph("________________________________________________________________________________________________________", style='HorizontalLine')
    
    education_details_paragraph = document.add_paragraph()
    education_details_paragraph.add_run(f"{cv_data['education'][0]['university']} - {cv_data['education'][0]['location']}")

    # Set tab stops for right alignment of start date and end date
    tab_stops = education_details_paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)  # Right align the tab at 6.5 inches

    # Add the start date and end date with right-aligned tab stops
    education_details_paragraph.add_run("\t")  # Add a tab character for right alignment
    education_details_paragraph.add_run(f"{cv_data['education'][0]['startDate']} - {cv_data['education'][0]['endDate']}")

    degree_gpa_paragraph = document.add_paragraph()
    degree_gpa_line = f"{cv_data['education'][0]['degree']} in {cv_data['education'][0]['major']} | GPA: {cv_data['education'][0]['gpa']:.2f}/5.00"
    degree_gpa_paragraph.add_run(degree_gpa_line)







    document.add_paragraph(plain_text, style='BodyText')

    # Save the Word document
    document.save(output_file)
# Rendering the template with the CV data
rendered_cv = template.render(cv_data=mockCVData[0])

# Convert HTML to Word document
output_file_path = 'output_cv19.docx'
create_word_document(mockCVData[0],rendered_cv, output_file_path)

print(f"Word document '{output_file_path}' created successfully.")